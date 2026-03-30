from typing import List, Dict, Any, Optional, Tuple, Union, Callable
from sqlalchemy.orm import Session
from ..database.models.hierarchy import Chapter, Section, Subsection, RawMaterial
from ..database.models.question import Question
from ..database.models.chunk import Chunk, KnowledgeRelation
from ..database.models.course import Course, IngestionStatus
from ..rag.embedder import Embedder
import os
import magic
from ..quiz.llm_service import llm


class MaterialProcessor:
    def __init__(self, db: Session):
        self.db = db

    # ------------------------------------------------------------------
    #  STATUS HELPERS
    # ------------------------------------------------------------------

    def _set_status(self, document, status: str):
        """Update ingestion status and commit so the professor sees live progress."""
        document.ingestion_status = status
        self.db.commit()

    def _fail(self, document, message: str):
        """Mark Document as FAILED with a human-readable error message."""
        from ..database.models.course import IngestionStatus
        document.ingestion_status = IngestionStatus.FAILED
        document.ingestion_error = message
        self.db.commit()

    # ------------------------------------------------------------------
    #  MAIN ENTRY POINT
    # ------------------------------------------------------------------

    def process_material(self, course_id: int, document_id: int, file_path: str, file_type: str):
        """
        Main entry point for processing a study material.
        Stages: VALIDATING → EXTRACTING → (OCR_PROCESSING) → CHUNKING → EMBEDDING → COMPLETED
        """
        import time
        from ..database.models.course import Document
        start_time = time.time()
        print(f"\n{'#'*60}")
        print(f"### [INGESTION ENGINE] Processing: {os.path.basename(file_path)}")
        print(f"{'#'*60}")

        document = self.db.query(Document).get(document_id)
        if not document:
            print(f"[!] INGESTION ABORTED: Document {document_id} not found.")
            return

        # Clear any previous error
        document.ingestion_error = None

        try:
            # ── STAGE 1: VALIDATION ──
            self._set_status(document, IngestionStatus.VALIDATING)
            validation_error = self._validate_file(file_path)
            if validation_error:
                self._fail(document, validation_error)
                print(f"[!] VALIDATION FAILED: {validation_error}")
                return

            # ── STAGE 2: MIME DETECTION & EXTRACTION ──
            self._set_status(document, IngestionStatus.EXTRACTING)

            # NOTE: Removed `self.clear_course_data(course_id)` to allow multiple documents
            # to be ingested into the same course simultaneously.

            mime_type = self._detect_mime_type(file_path)
            print(f"[*] Detected MIME type: {mime_type}")

            extracted_text: Optional[str]
            extracted_data: Optional[List[Dict[str, Any]]]
            extracted_text, extracted_data = self._dispatch_extraction(
                file_path, mime_type, document
            )

            if extracted_data is None and not extracted_text:
                self._fail(document, f"No content could be extracted from the file.")
                return

            # If we got flat text but no hierarchy, wrap it into a default structure
            if extracted_data is None and extracted_text:
                extracted_data = self._text_to_hierarchy(extracted_text, file_path)

            if not extracted_data:
                self._fail(document, "Extraction produced no usable content.")
                return

            # ── STAGE 2.5: AUTHOR EXTRACTION ──
            try:
                # Use the first 2000 characters of extracted text or first chunk
                sample_text = ""
                if extracted_text:
                    sample_text = extracted_text[:4000]
                elif extracted_data:
                    # Flatten first few chunks
                    sample_text = "\n".join([
                        s["content"] for c in extracted_data[:1] 
                        for sec in c["sections"] 
                        for s in sec["subsections"]
                    ])[:4000]
                
                if sample_text:
                    document.author = self._extract_author(sample_text)
                    print(f"[*] Identified document author: {document.author}")
            except Exception as ae:
                print(f"[!] Author extraction failed (non-fatal): {ae}")
                document.author = "the author"

            # ── STAGE 3: CHUNKING + EMBEDDING (inside _store_hierarchy) ──
            self._set_status(document, IngestionStatus.CHUNKING)
            self._store_hierarchy(course_id, extracted_data, document)

            # ── DONE ──
            duration = time.time() - start_time
            print(f"\n{'='*60}")
            print(f"✅ [SUCCESS] Material chunked and indexed in "
                  f"{duration:.2f}s")
            print(f"⚙️  Concept extraction running in background...")
            print(f"🔗 Quiz is available now. Adaptive mode activates "
                  f"once concept extraction completes.")
            print(f"{'='*60}\n")

        except Exception as e:
            self.db.rollback()
            from ..database.models.course import Document
            document = self.db.query(Document).get(document_id)
            if document:
                self._fail(document, str(e))
            print(f"\n❌ [FATAL ERROR] Ingestion Pipeline Failed: {e}")

    # ------------------------------------------------------------------
    #  STAGE 1: PRE-INGESTION VALIDATION
    # ------------------------------------------------------------------

    def _validate_file(self, file_path: str) -> Optional[str]:
        """
        Validates the file before any extraction.
        Returns an error message string, or None if valid.
        """
        # Check exists
        if not os.path.exists(file_path):
            return f"File not found: {os.path.basename(file_path)}"

        # Check size > 0
        if os.path.getsize(file_path) == 0:
            return f"File is empty: {os.path.basename(file_path)}"

        # Check for corrupt/password-protected PDFs
        if file_path.lower().endswith(".pdf"):
            try:
                import fitz
                doc = fitz.open(file_path)
                if doc.is_encrypted:
                    authenticated = doc.authenticate("")
                    if not authenticated:
                        doc.close()
                        return "PDF is password-protected. Please upload an unprotected file."
                doc.close()
            except Exception as e:
                return f"File appears to be corrupt or unreadable: {e}"

        # Check for corrupt DOCX
        if file_path.lower().endswith((".docx", ".doc")):
            try:
                import docx
                docx.Document(file_path)
            except Exception as e:
                return f"DOCX file appears to be corrupt: {e}"

        return None

    # ------------------------------------------------------------------
    #  STAGE 2: MIME DETECTION & DISPATCH
    # ------------------------------------------------------------------

    def _detect_mime_type(self, file_path: str) -> str:
        """Detect actual MIME type using python-magic, ignoring file extension."""
        try:
            mime = magic.from_file(file_path, mime=True)
            return mime
        except Exception as e:
            print(f"[!] MIME detection failed: {e}, falling back to extension")
            ext = file_path.lower().rsplit(".", 1)[-1] if "." in file_path else ""
            fallback_map = {
                "pdf": "application/pdf",
                "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                "txt": "text/plain",
                "png": "image/png",
                "jpg": "image/jpeg",
                "jpeg": "image/jpeg",
                "tiff": "image/tiff",
                "tif": "image/tiff",
            }
            return fallback_map.get(ext, "application/octet-stream")

    def _dispatch_extraction(
        self, file_path: str, mime_type: str, document
    ) -> Tuple[Optional[str], Optional[List[Dict[str, Any]]]]:
        """
        Routes to the correct extraction handler based on MIME type.
        Returns (flat_text, hierarchy_data) — one or both may be None.
        """
        MIME_HANDLERS: Dict[str, Callable[..., Any]] = {
            "application/pdf": self._extract_pdf,
            "application/vnd.openxmlformats-officedocument.wordprocessingml.document": self._extract_docx,
            "text/plain": self._extract_txt,
            "image/png": self._extract_image,
            "image/jpeg": self._extract_image,
            "image/tiff": self._extract_image,
        }

        handler = MIME_HANDLERS.get(mime_type)
        if handler is None:
            self._fail(document, f"Unsupported file type: {mime_type}")
            return None, None

        # Image and TXT handlers return flat text
        if mime_type.startswith("image/") or mime_type == "text/plain":
            self._set_status(document, IngestionStatus.OCR_PROCESSING if mime_type.startswith("image/") else IngestionStatus.EXTRACTING)
            text: str = handler(file_path, document)
            return text, None

        # PDF/DOCX handlers return hierarchy data
        hierarchy: List[Dict[str, Any]] = handler(file_path, document)
        return None, hierarchy

    # ------------------------------------------------------------------
    #  PDF EXTRACTION (with per-page OCR for scanned pages)
    # ------------------------------------------------------------------

    def _extract_pdf(self, file_path: str, document) -> List[Dict[str, Any]]:
        """
        Improved PDF extraction with per-page scanned detection and OCR.
        """
        import fitz

        print(f"[*] Audit Phase 1: Deep Text Extraction")

        try:
            doc = fitz.open(file_path)
        except Exception as e:
            self._fail(document, f"Failed to open PDF: {e}")
            return []

        if doc.is_encrypted and not doc.authenticate(""):
            doc.close()
            self._fail(document, "PDF is password-protected")
            return []

        total_pages = len(doc)
        print(f"    -> Pages Detected: {total_pages}")

        page_texts = []
        needs_ocr = False

        for i, page in enumerate(doc):
            page_text = page.get_text("text").replace("\0", "")

            # Readability ratio: text density relative to page area
            page_area = max(page.rect.width * page.rect.height / 100, 1)
            readability = len(page_text.strip()) / page_area

            if readability < 0.8:
                # Page is likely scanned
                needs_ocr = True
            else:
                page_texts.append((i, page_text))

            if total_pages < 20 or (i + 1) % 10 == 0 or (i + 1) == total_pages:
                status = "SCANNED" if readability < 0.8 else "OK"
                print(f"    [PROGRESS] Page {i+1}/{total_pages} extracted. [{status}]")

        # OCR pass for scanned pages
        if needs_ocr:
            self._set_status(document, IngestionStatus.OCR_PROCESSING)
            print(f"[*] Running OCR on scanned pages...")

            import pytesseract
            from PIL import Image
            import io

            scanned_indices = set(range(total_pages)) - set(idx for idx, _ in page_texts)

            for i in sorted(scanned_indices):
                page = doc[i]
                pix = page.get_pixmap(dpi=200)
                img = Image.open(io.BytesIO(pix.tobytes("png")))
                ocr_text = pytesseract.image_to_string(img)
                page_texts.append((i, ocr_text))
                print(f"    [OCR] Page {i+1}/{total_pages} → {len(ocr_text)} chars extracted")

        doc.close()

        # Sort pages back into order and join
        page_texts.sort(key=lambda x: x[0])
        full_text = "\n\n".join(text for _, text in page_texts)

        if not full_text.strip():
            return []

        print(f"[*] Audit Phase 2: Hierarchical Syllabus Mapping")

        # Build chapter hierarchy (same grouping logic as before)
        chapters = []
        pages_per_chapter = 5

        doc = fitz.open(file_path)
        for i in range(0, total_pages, pages_per_chapter):
            end_page = min(i + pages_per_chapter, total_pages)

            # Use our already-extracted text for these pages
            chapter_text = "\n\n".join(
                text for idx, text in page_texts if i <= idx < end_page
            )

            chap_num = (i // pages_per_chapter) + 1
            chapters.append({
                "title": f"Chapter {chap_num} (Pages {i+1}-{end_page})",
                "order": chap_num,
                "sections": [{
                    "title": f"Section {chap_num}.1",
                    "order": 1,
                    "subsections": [{
                        "title": f"Content Block {chap_num}.1.1",
                        "order": 1,
                        "content": chapter_text
                    }]
                }]
            })
        doc.close()

        print(f"    -> ACCURACY UPGRADE: Mapped {total_pages} pages into {len(chapters)} logical Chapters.")
        return chapters

    # ------------------------------------------------------------------
    #  DOCX EXTRACTION (heading-aware)
    # ------------------------------------------------------------------

    def _extract_docx(self, file_path: str, document) -> List[Dict[str, Any]]:
        """
        Improved DOCX extraction. If headings are found, uses them for structure.
        Falls back to flat paragraph joining if no headings detected.
        """
        import docx as python_docx

        doc = python_docx.Document(file_path)

        # Check for headings
        has_headings = False
        structured = []
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                has_headings = True
                try:
                    level = int(style_name.split()[-1])
                except (ValueError, IndexError):
                    level = 1
                structured.append({"level": level, "text": text})
            else:
                structured.append({"level": 0, "text": text})

        if has_headings:
            # Build hierarchy from headings
            return self._build_hierarchy_from_headings(structured)
        else:
            # Fallback: join all paragraphs, group into chapters
            full_text = "\n".join(item["text"] for item in structured)
            return self._text_to_hierarchy(full_text, file_path)

    def _build_hierarchy_from_headings(self, structured: List[Dict[str, Any]]) -> List[Dict[str, Any]]:
        """Converts heading-structured list into chapter/section/subsection hierarchy."""
        chapters: List[Dict[str, Any]] = []
        current_chapter: Optional[Dict[str, Any]] = None
        current_section: Optional[Dict[str, Any]] = None
        body_buffer: List[str] = []

        def flush_body(target_sub_list: List[Dict[str, Any]], title: str) -> None:
            """Saves accumulated body text as a subsection."""
            if body_buffer:
                content = "\n".join(body_buffer)
                target_sub_list.append({
                    "title": title,
                    "order": len(target_sub_list) + 1,
                    "content": content
                })
                body_buffer.clear()

        for item in structured:
            level = item["level"]
            text = item["text"]

            if level == 1:
                # New chapter
                if current_section and current_chapter:
                    flush_body(current_section["subsections"], f"Content")
                if current_chapter:
                    if not current_chapter["sections"]:
                        current_chapter["sections"].append({
                            "title": "Main Section",
                            "order": 1,
                            "subsections": []
                        })
                        current_section = current_chapter["sections"][-1]
                    flush_body(current_section["subsections"], "Content")
                    chapters.append(current_chapter)

                current_chapter = {
                    "title": text,
                    "order": len(chapters) + 1,
                    "sections": []
                }
                current_section = None

            elif level == 2:
                # New section
                if current_chapter is None:
                    current_chapter = {
                        "title": "Document",
                        "order": 1,
                        "sections": []
                    }
                if current_section:
                    flush_body(current_section["subsections"], "Content")

                current_section = {
                    "title": text,
                    "order": len(current_chapter["sections"]) + 1,
                    "subsections": []
                }
                current_chapter["sections"].append(current_section)

            elif level == 3:
                # Subsection heading — flush body, start new subsection
                if current_chapter is None:
                    current_chapter = {"title": "Document", "order": 1, "sections": []}
                if current_section is None:
                    current_section = {"title": "Main Section", "order": 1, "subsections": []}
                    current_chapter["sections"].append(current_section)
                flush_body(current_section["subsections"], "Content")
                # Next body text will go under this heading
                body_buffer.append(text)

            else:
                # Body paragraph
                body_buffer.append(text)

        # Flush remaining
        if current_chapter is None:
            current_chapter = {"title": "Document", "order": 1, "sections": []}
        if current_section is None:
            current_section = {"title": "Main Section", "order": 1, "subsections": []}
            current_chapter["sections"].append(current_section)
        flush_body(current_section["subsections"], "Content")
        chapters.append(current_chapter)

        return chapters

    # ------------------------------------------------------------------
    #  TXT EXTRACTION
    # ------------------------------------------------------------------

    def _extract_txt(self, file_path: str, document) -> str:
        """Reads a plain text file and returns its content."""
        print(f"[*] Extracting plain text from: {os.path.basename(file_path)}")
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        print(f"    -> Extracted {len(text)} characters")
        return text

    # ------------------------------------------------------------------
    #  IMAGE OCR EXTRACTION
    # ------------------------------------------------------------------

    def _extract_image(self, file_path: str, document) -> str:
        """
        Extracts text from an image using pytesseract OCR.
        Preprocesses: grayscale + sharpen.
        """
        import pytesseract
        from PIL import Image, ImageFilter

        print(f"[*] Image OCR: {os.path.basename(file_path)}")
        self._set_status(document, IngestionStatus.OCR_PROCESSING)

        img = Image.open(file_path)
        img = img.convert("L")  # Grayscale
        img = img.filter(ImageFilter.SHARPEN)

        text = pytesseract.image_to_string(img)
        print(f"    -> OCR extracted {len(text)} characters")

        if not text.strip():
            self._fail(document, "OCR could not extract any text from the image.")
            return ""

        return text

    # ------------------------------------------------------------------
    #  TEXT → HIERARCHY HELPER
    # ------------------------------------------------------------------

    def _text_to_hierarchy(self, text: str, file_path: str) -> List[Dict[str, Any]]:
        """Wraps flat text into a default chapter/section/subsection structure."""
        chunk_size = 2000
        text_chunks = [text[i:i+chunk_size] for i in range(0, len(text), chunk_size)]

        chapters = []
        pages_per_chapter = 5

        for i in range(0, len(text_chunks), pages_per_chapter):
            end_chunk = min(i + pages_per_chapter, len(text_chunks))
            chapter_text = "".join(text_chunks[i:end_chunk])
            chap_num = (i // pages_per_chapter) + 1
            chapters.append({
                "title": f"Chapter {chap_num} (Parts {i+1}-{end_chunk})",
                "order": chap_num,
                "sections": [{
                    "title": f"Section {chap_num}.1",
                    "order": 1,
                    "subsections": [{
                        "title": f"Content Block {chap_num}.1.1",
                        "order": 1,
                        "content": chapter_text
                    }]
                }]
            })
        return chapters

    # ------------------------------------------------------------------
    #  STAGE 2.5: AUTHOR EXTRACTION (LLM)
    # ------------------------------------------------------------------

    def _extract_author(self, text: str) -> str:
        """Uses LLM to identify the main author of the text."""
        system_prompt = (
            "You are a metadata extraction assistant. "
            "Identify the main author's name from the provided text excerpt. "
            "Respond ONLY with the author's last name or full name. "
            "If not found, respond with 'the author'."
        )
        prompt = f"Identify the primary author of this text:\n\n{text[:2000]}"
        
        try:
            # Using fast mode (Hugging Face) for metadata tasks
            result = llm.generate_content_fast(prompt, system_prompt=system_prompt)
            if "ERROR" in result or len(result) > 50:
                return "the author"
            return result.strip()
        except Exception:
            return "the author"

    # ------------------------------------------------------------------
    #  STORE HIERARCHY (preserved from original — only status tracking added)
    # ------------------------------------------------------------------

    def _store_hierarchy(self, course_id: int, hierarchy_data: List[Dict[str, Any]], document=None):
        """Saves the detected hierarchy to the database and triggers RAG updates."""
        from .chunking import Chunker
        from ..rag.embedder import Embedder

        _concept_extraction_queue = []

        chunker = Chunker(self.db)
        embedder = Embedder(self.db)

        print(f"  > Storing {len(hierarchy_data)} chapters to DB...")

        for chap_data in hierarchy_data:
            chapter = Chapter(
                title=chap_data["title"], 
                order=chap_data["order"], 
                course_id=course_id,
                document_id=document.id if document else None
            )
            self.db.add(chapter)
            self.db.flush()

            for sec_data in chap_data["sections"]:
                section = Section(title=sec_data["title"], order=sec_data["order"], chapter_id=chapter.id)
                self.db.add(section)
                self.db.flush()

                for sub_data in sec_data["subsections"]:
                    try:
                        subsection = Subsection(title=sub_data["title"], order=sub_data["order"], section_id=section.id)
                        self.db.add(subsection)
                        self.db.flush()

                        raw_mat = RawMaterial(content=sub_data["content"], subsection_id=subsection.id)
                        self.db.add(raw_mat)
                        self.db.flush()

                        print(f"    - Processing {subsection.title}...")

                        # CHUNKING stage
                        if document:
                            self._set_status(document, IngestionStatus.CHUNKING)
                        chunker.generate_chunks(subsection.id, document_id=document.id if document else None)

                        # EMBEDDING stage
                        print(f"    - Indexing {subsection.title} in pgvector...")
                        if document:
                            self._set_status(document, IngestionStatus.EMBEDDING)
                        try:
                            embedder.embed_chunks(subsection.id)
                        except Exception as ee:
                            print(f"      [EMBEDDING WARNING] {ee}")

                        self.db.commit()  # Persistent save for each subsection

                        # Concept extraction runs in background after full ingestion
                        # collect subsection IDs for background phase
                        _concept_extraction_queue.append(subsection.id)

                    except Exception as sub_e:
                        print(f"    [SUBSECTION ERROR] {sub_e}")
                        self.db.rollback()
                        raise sub_e

        # Signal that basic ingestion is complete
        if document:
            self._set_status(document, IngestionStatus.COMPLETED)
            self.db.commit()

        # Run concept extraction in background thread
        if _concept_extraction_queue and course_id:
            import threading
            t = threading.Thread(
                target=self._run_concept_extraction_background,
                args=(course_id, _concept_extraction_queue, document.id if document else None),
                daemon=True
            )
            t.start()
            print(f"[Processor] Concept extraction started in background "
                  f"for {len(_concept_extraction_queue)} subsections")

    def _run_concept_extraction_background(
        self,
        course_id: int,
        subsection_ids: List[int],
        document_id: int
    ):
        """
        Runs concept extraction in a background thread.
        Uses a FRESH database session — never shares session
        with the main ingestion thread.
        Sets Document status to CONCEPT_EXTRACTION while running,
        FULLY_READY when done, FAILED if it errors.
        """
        from ..database.session import SessionLocal
        from ..database.models.course import Document, IngestionStatus
        from .concept_extractor import ConceptExtractor

        db = SessionLocal()
        try:
            # Update status to CONCEPT_EXTRACTION
            document = db.query(Document).get(document_id)
            if document:
                document.ingestion_status = IngestionStatus.CONCEPT_EXTRACTION
                db.commit()

            extractor = ConceptExtractor(db)
            total = len(subsection_ids)

            for i, subsection_id in enumerate(subsection_ids, 1):
                try:
                    print(f"[ConceptBG] Subsection {i}/{total} "
                          f"(id={subsection_id})")
                    extractor.extract_and_store(subsection_id, course_id)
                except Exception as e:
                    print(f"[ConceptBG] Non-fatal error on "
                          f"subsection {subsection_id}: {e}")
                    continue  # Skip failed subsection, keep going

            # All done
            document = db.query(Document).get(document_id)
            if document:
                document.ingestion_status = IngestionStatus.FULLY_READY
                db.commit()
            print(f"[ConceptBG] Concept extraction complete "
                  f"for Document {document_id}")

        except Exception as e:
            print(f"[ConceptBG] Fatal background error: {e}")
            try:
                document = db.query(Document).get(document_id)
                if document:
                    document.ingestion_status = IngestionStatus.FAILED
                    document.ingestion_error = (
                        f"Concept extraction failed: {str(e)}"
                    )
                    db.commit()
            except Exception:
                pass
        finally:
            db.close()

    # ------------------------------------------------------------------
    #  CLEAR COURSE DATA (preserved from original)
    # ------------------------------------------------------------------

    def clear_course_data(self, course_id: int):
        """Wipes all hierarchical and assessment data for a course to prevent leakage."""
        print(f"\n[*] CLEANUP: Wiping stale data for Course {course_id}...")

        # 1. Reset embeddings
        embedder = Embedder(self.db)
        embedder.reset_index()

        # 2. Clear DB (Hierarchical cascade deletes Sections, Subsections, RawMaterial, and Chunks)
        chapters = self.db.query(Chapter).filter_by(course_id=course_id).all()
        for chapter in chapters:
            self.db.query(Question).filter(
                Question.subsection_id.in_(
                    self.db.query(Subsection.id).join(Section).filter(Section.chapter_id == chapter.id)
                )
            ).delete(synchronize_session=False)

            self.db.delete(chapter)

        # 3. Clear concepts for this course
        from ..database.models.concept import Concept
        self.db.query(Concept).filter_by(course_id=course_id).delete(synchronize_session=False)

        self.db.commit()
        print("    -> Database cleared.")

    # ------------------------------------------------------------------
    #  KNOWLEDGE GRAPH
    # ------------------------------------------------------------------

    def _create_deterministic_relations(self, subsection_id: int):
        """
        DEPRECATED: replaced by LLM concept extraction in Phase 4.
        Builds KnowledgeRelations by matching keywords between the new subsection and existing ones.
        """
        from ..database.models.chunk import Chunk, KnowledgeRelation, ChunkType
        from ..database.models.hierarchy import Subsection
        import re

        current_chunks = self.db.query(Chunk).filter_by(subsection_id=subsection_id, chunk_type=ChunkType.MEDIUM).all()
        if not current_chunks:
            return

        ignore_words = {"this", "that", "there", "their", "chapter", "section", "about", "would", "could", "should"}
        keywords = set()
        for chunk in current_chunks:
            found = re.findall(r"\b[A-Z][a-z]{4,}\b", chunk.content)
            for word in found:
                if word.lower() not in ignore_words:
                    keywords.add(word)

        if not keywords:
            return

        current_sub = self.db.query(Subsection).get(subsection_id)
        if not current_sub:
            return

        other_chunks = self.db.query(Chunk).join(Subsection).join(Section).join(Chapter).filter(
            Chapter.course_id == current_sub.section.chapter.course_id,
            Subsection.id != subsection_id,
            Chunk.chunk_type == ChunkType.MEDIUM
        ).all()

        for cur_chunk in current_chunks:
            for other_chunk in other_chunks:
                shared = [k for k in keywords if k in other_chunk.content]
                if shared:
                    relation = KnowledgeRelation(
                        source_id=cur_chunk.id,
                        target_id=other_chunk.id,
                        relation_type=f"shared_concept:{shared[0]}"
                    )
                    self.db.add(relation)

        self.db.commit()

    def _create_semantic_relations(self, subsection_id: int):
        """[DEPRECATED] AI-based relation builder - preserved for compatibility check."""
        pass
